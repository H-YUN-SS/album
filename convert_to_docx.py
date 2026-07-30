import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def convert_txt_to_docx(input_file, output_file):
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(11)

    # 读取文件
    with open(input_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_code_block = False
    code_lines = []
    in_table = False
    table_lines = []

    for line in lines:
        line = line.rstrip('\n')

        # 跳过纯分隔线
        if re.match(r'^={10,}$', line):
            continue
        if re.match(r'^-{10,}$', line):
            continue

        # 处理标题行（带🔴🟡🟢⚪的部分标题）
        if re.match(r'^[🔴🟡🟢⚪]\s*第.+部分', line):
            p = doc.add_heading(line.strip(), level=1)
            continue

        # 处理Q:开头的问题
        if line.startswith('Q:'):
            p = doc.add_paragraph()
            run = p.add_run(line.strip())
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 102, 204)
            continue

        # 处理A:开头的答案标记
        if line.strip() == 'A:':
            p = doc.add_paragraph()
            run = p.add_run('答：')
            run.bold = True
            continue

        # 处理表格行
        if '│' in line or '├' in line or '└' in line or '┌' in line:
            # 清理表格边框字符
            clean_line = line.replace('┌', '').replace('┐', '').replace('└', '').replace('┘', '')
            clean_line = clean_line.replace('├', '').replace('┤', '').replace('│', ' | ')
            clean_line = clean_line.replace('─', '-')
            clean_line = clean_line.strip()
            if clean_line:
                p = doc.add_paragraph(clean_line)
                p.style = doc.styles['Normal']
                pf = p.paragraph_format
                pf.space_before = Pt(0)
                pf.space_after = Pt(0)
            continue

        # 处理代码块
        if line.strip().startswith('//') or line.strip().startswith('class ') or \
           line.strip().startswith('void ') or line.strip().startswith('int ') or \
           line.strip().startswith('if ') or line.strip().startswith('for ') or \
           line.strip().startswith('while ') or line.strip().startswith('{') or \
           line.strip().startswith('}') or line.strip().startswith('public:') or \
           line.strip().startswith('private:') or line.strip().startswith('protected:'):
            if not in_code_block:
                in_code_block = True
                code_lines = []
            code_lines.append(line)
            continue

        # 如果之前在代码块中，现在不在了，输出代码块
        if in_code_block and line.strip():
            in_code_block = False
            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph()
            run = p.add_run(code_text)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0, 0, 0)
            p.paragraph_format.left_indent = Inches(0.3)
            code_lines = []

        # 处理带emoji的标记行
        if re.match(r'^[🔴🟡🟢⚪]\s*【', line):
            p = doc.add_paragraph(line.strip())
            p.runs[0].bold = True if '🔴' in line or '🟡' in line else False
            continue

        # 处理普通内容行
        if line.strip():
            # 检查是否是列表项
            if line.strip().startswith('- ') or line.strip().startswith('• '):
                p = doc.add_paragraph(line.strip(), style='List Bullet')
            elif re.match(r'^\d+[\.\)]\s', line.strip()):
                p = doc.add_paragraph(line.strip(), style='List Number')
            else:
                p = doc.add_paragraph(line.strip())

    # 处理最后的代码块
    if in_code_block and code_lines:
        code_text = '\n'.join(code_lines)
        p = doc.add_paragraph()
        run = p.add_run(code_text)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        p.paragraph_format.left_indent = Inches(0.3)

    doc.save(output_file)
    print(f"转换完成！输出文件：{output_file}")

if __name__ == '__main__':
    input_file = r'D:\xwechat_files\wxid_4wjn2dflpatn22_3223\msg\file\2026-07\八股_完整版.txt'
    output_file = r'D:\xwechat_files\wxid_4wjn2dflpatn22_3223\msg\file\2026-07\八股_完整版.docx'
    convert_txt_to_docx(input_file, output_file)
